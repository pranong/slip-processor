from utils.logger import log
"""
gen_pdf.py — gen PDF ใบรับรองแทนใบเสร็จรับเงิน จาก template .docx
- อ่าน slip_data JSON ที่ sort_slips.py สร้างไว้ (ไม่ call Claude API ซ้ำ)
- clone row ในตาราง + เติมข้อมูล
- gen 1 PDF ต่อ batch ใหม่แต่ละวัน
- merge summary.json
"""

import json
import shutil
import subprocess
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from config.config import DATA_MOUNT, TEMPLATE_DIR, TEMPLATE_PATH, MONTH_MAP
from config.gen_config import (
    NOTE_ROUTES, NOTE_DEFAULT_SUBFOLDER,
    NOTE_DEFAULT_CERT_TEMPLATE, NOTE_DEFAULT_RECEIPT_TEMPLATE,
)
from utils.thai_baht_text import baht_text
from utils.state import load_state, save_state, mark_generated, is_generated
from utils.vendor import find_vendor, load_vendors

MONTH_MAP_NUM = {
    "JAN": "01", "FEB": "02", "MAR": "03", "APR": "04",
    "MAY": "05", "JUN": "06", "JUL": "07", "AUG": "08",
    "SEP": "09", "OCT": "10", "NOV": "11", "DEC": "12",
}

# ── Routing ───────────────────────────────────────────────────────────────────
# ย้ายไปอยู่ที่ config.py แล้ว — เพิ่ม/ลด keyword ได้ที่นั่น


def get_route(note: str) -> dict:
    """หา subfolder + template (cert และ receipt) จาก note — เจอ keyword ที่ไหนก็ได้ ไม่ case sensitive"""
    note_lower = (note or "").lower()
    for route in NOTE_ROUTES:
        if route["keyword"].lower() in note_lower:
            return {
                "subfolder":        route["subfolder"],
                "cert_template":    Path(TEMPLATE_DIR) / route["cert_template"],
                "receipt_template": Path(TEMPLATE_DIR) / route["receipt_template"],
            }
    return {
        "subfolder":        NOTE_DEFAULT_SUBFOLDER,
        "cert_template":    Path(TEMPLATE_DIR) / NOTE_DEFAULT_CERT_TEMPLATE,
        "receipt_template": Path(TEMPLATE_DIR) / NOTE_DEFAULT_RECEIPT_TEMPLATE,
    }

# ── Summary ───────────────────────────────────────────────────────────────────

def _summary_path(year_dir: Path) -> Path:
    return year_dir / "summary.json"


def load_summary(year_dir: Path) -> dict:
    p = _summary_path(year_dir)
    if p.exists():
        return json.loads(p.read_text(encoding="utf-8"))
    return {}


def save_summary(year_dir: Path, summary: dict):
    # backup ก่อน
    p = _summary_path(year_dir)
    if p.exists():
        bak_dir = year_dir / "backups"
        bak_dir.mkdir(exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy(p, bak_dir / f"summary_{ts}.json")

    p.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")


def merge_to_summary(summary: dict, month_key: str, day_str: str,
                     slips: list[dict], pdf_name: str):
    """เพิ่ม transaction เข้า summary แบบ merge"""
    if month_key not in summary:
        summary[month_key] = {"total": 0, "count": 0, "transactions": []}

    for s in slips:
        amt = s.get("amount") or 0
        summary[month_key]["total"] += amt
        summary[month_key]["count"] += 1
        summary[month_key]["transactions"].append({
            "date":        f"{s.get('day',''):02d}/{month_key}/{s.get('year_ce','')}",
            "amount":      amt,
            "description": s.get("description"),
            "from":        s.get("from_account"),
            "to":          s.get("to_account"),
            "bank":        s.get("bank"),
            "ref":         s.get("ref"),
            "pdf":         pdf_name,
        })

# ── DOCX template fill ────────────────────────────────────────────────────────

def _clone_row(table, row_idx: int):
    """clone row ในตาราง (XML level) แล้ว insert หลัง row เดิม"""
    src_tr = table.rows[row_idx]._tr
    new_tr = deepcopy(src_tr)
    src_tr.addnext(new_tr)
    return new_tr


def _set_cell_text(row_element, col_idx: int, text: str):
    """เขียนข้อความลง cell โดยรักษา format เดิม"""
    cells = row_element.findall(qn("w:tc"))
    if col_idx >= len(cells):
        return
    cell = cells[col_idx]
    # หา paragraph แรก
    para = cell.find(qn("w:p"))
    if para is None:
        return
    # หา run แรก (ถ้ามี) เพื่อเอา format
    runs = para.findall(qn("w:r"))
    if runs:
        # ลบ run เดิมทั้งหมด
        for r in runs:
            para.remove(r)
        # สร้าง run ใหม่ด้วย format เดิม
        new_run = deepcopy(runs[0])
        # เปลี่ยนข้อความ
        t_elem = new_run.find(qn("w:t"))
        if t_elem is not None:
            t_elem.text = text
            t_elem.set(qn("xml:space"), "preserve")
        para.append(new_run)
    else:
        # ไม่มี run — สร้างใหม่
        from docx.oxml import OxmlElement
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        run.append(t)
        para.append(run)


def _replace_across_runs(paragraph, placeholder: str, value: str):
    """แทน placeholder ที่อาจกระจายหลาย runs — merge เฉพาะ runs ที่มี placeholder"""
    runs = paragraph.runs
    if not runs:
        return
    full_text = "".join(r.text for r in runs)
    if placeholder not in full_text:
        return

    # หาว่า placeholder เริ่มและจบที่ run ไหน
    pos = full_text.index(placeholder)
    end_pos = pos + len(placeholder)

    char_count = 0
    start_run = None
    end_run = None
    for i, r in enumerate(runs):
        run_start = char_count
        run_end = char_count + len(r.text)
        if start_run is None and run_end > pos:
            start_run = i
        if run_end >= end_pos:
            end_run = i
            break
        char_count = run_end

    if start_run is None or end_run is None:
        return

    # รวมข้อความเฉพาะ runs ที่มี placeholder
    merged = "".join(runs[i].text for i in range(start_run, end_run + 1))
    merged = merged.replace(placeholder, value)

    # ใส่กลับใน run แรกของกลุ่ม ลบ run ที่เหลือในกลุ่ม
    runs[start_run].text = merged
    for i in range(start_run + 1, end_run + 1):
        runs[i].text = ""


def _replace_text_in_doc(doc: DocxDocument, placeholder: str, value: str):
    """แทน placeholder ทั่วทั้ง document (paragraphs + table cells)"""
    for para in doc.paragraphs:
        if placeholder in para.text:
            _replace_across_runs(para, placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        _replace_across_runs(para, placeholder, value)


def fill_template(slips: list[dict], day_str: str, month_name: str,
                  year_thai: str, template_path: Path | None = None) -> Path | None:
    """เติมข้อมูลลง template, return path ของ docx ที่ fill แล้ว"""
    template = template_path if template_path else Path(TEMPLATE_PATH)
    if not template.exists():
        log(f"    ⚠️  ไม่พบ template: {template}")
        return None

    doc = DocxDocument(str(template))

    # หาตารางแรก
    if not doc.tables:
        log("    ⚠️  ไม่พบตารางใน template")
        return None
    table = doc.tables[0]

    # หา template row — รองรับทั้ง format ใหม่ ($seq[0]) และเก่า ($date[0])
    template_row_idx = None
    total_row_idx = None
    is_new_format = False  # True = ใบรับรองฯ ใหม่ ($seq), False = บริษัท ($date[0])
    for idx, row in enumerate(table.rows):
        row_text = "".join(cell.text for cell in row.cells)
        if "$seq[0]" in row_text:
            template_row_idx = idx
            is_new_format = True
        elif "$date[0]" in row_text and template_row_idx is None:
            template_row_idx = idx
            is_new_format = False
        if "$intSumTotal" in row_text:
            total_row_idx = idx

    if template_row_idx is None:
        log("    ⚠️  ไม่พบ $seq[0] หรือ $date[0] ใน template")
        return None

    # ── clone row ถ้ามีมากกว่า 1 slip ──
    if len(slips) > 1:
        for _ in range(len(slips) - 1):
            _clone_row(table, template_row_idx)

    # ── เติมข้อมูลแต่ละ row ──
    total = 0
    for i, slip in enumerate(slips):
        row_idx = template_row_idx + i
        tr = table.rows[row_idx]._tr

        desc_str = slip.get("note") or slip.get("description") or "-"
        amt      = slip.get("amount") or 0
        amt_str  = f"{amt:,.2f}"
        total   += amt

        if is_new_format:
            # format ใหม่: ลำดับ | รายละเอียด | จำนวนเงิน | หมายเหตุ
            _set_cell_text(tr, 0, str(i + 1))  # ลำดับ
            _set_cell_text(tr, 1, desc_str)
            _set_cell_text(tr, 2, amt_str)
        else:
            # format เก่า: วันที่ | รายละเอียด | จำนวนเงิน | หมายเหตุ
            day_val  = slip.get("day", "")
            mon_val  = slip.get("month", "")
            year_val = slip.get("year_ce", "")
            year_be  = year_val + 543 if isinstance(year_val, int) else ""
            date_str = f"{day_val}/{mon_val}/{year_be}" if day_val else ""
            _set_cell_text(tr, 0, date_str)
            _set_cell_text(tr, 1, desc_str)
            _set_cell_text(tr, 2, amt_str)

    # เติม total
    total_str    = f"{total:,.2f}"
    total_th_str = baht_text(total)

    # วันที่เอกสาร จาก slip แรก
    first  = slips[0] if slips else {}
    day_v  = first.get("day", "")
    mon_v  = first.get("month", "")
    yr_v   = first.get("year_ce", "")
    yr_be  = yr_v + 543 if isinstance(yr_v, int) else ""
    doc_date = f"{day_v}/{mon_v}/{yr_be}" if day_v else ""

    # แทน placeholders (ยาวก่อนสั้น กัน $date ไป match $fromDate)
    _replace_text_in_doc(doc, "$intSumTotal", total_str)
    _replace_text_in_doc(doc, "$thSumTotal", total_th_str)
    _replace_text_in_doc(doc, "$fromDate", doc_date)
    _replace_text_in_doc(doc, "$toDate", doc_date)
    if is_new_format:
        _replace_text_in_doc(doc, "$date", doc_date)

    # บันทึก docx ชั่วคราว
    import tempfile
    tmp = Path(tempfile.mkdtemp()) / f"filled_{month_name}_{day_str}.docx"
    doc.save(str(tmp))
    return tmp


def fill_template_receipt(slips: list[dict], day_str: str, month_name: str,
                          year_thai: str, template_path: Path | None = None) -> Path | None:
    """
    เติมข้อมูลลง template ใบสำคัญรับเงิน (รองรับทั้งตัวบุคคลและตัวบริษัท — หน้าตา/placeholder เหมือนกัน)
    placeholder ต่างจากใบรับรองฯ:
    - vendor info: $vndName, $vndId, $vndAddress, $vndM, $vndSt, $vndZone, $vndCity, $vndProv
    - ตาราง: $desc[0], $amountF[0], $amountSt[0]
    - รวม: $intSumFTotal, $intSumStTotal, $thSumTotal
    """
    template_path = template_path if template_path else Path(TEMPLATE_DIR) / "ใบสำคัญรับเงิน.docx"
    if not template_path.exists():
        log(f"    ⚠️  ไม่พบ template: {template_path}")
        return None

    doc = DocxDocument(str(template_path))

    if not doc.tables:
        log("    ⚠️  ไม่พบตารางใน template ใบสำคัญรับเงิน")
        return None

    # ── vendor info (จาก slip แรก — vendor เดียวกันทั้งวัน) ──
    first   = slips[0] if slips else {}
    vendor  = first.get("vendor") or {}
    day_v   = first.get("day", "")
    mon_v   = first.get("month", "")
    yr_v    = first.get("year_ce", "")
    yr_be   = yr_v + 543 if isinstance(yr_v, int) else ""
    doc_date = f"{day_v}/{mon_v}/{yr_be}" if day_v else ""

    # แทน vendor placeholders
    _replace_text_in_doc(doc, "$fromDate",   doc_date)
    _replace_text_in_doc(doc, "$vndName",    vendor.get("ชื่อ", "-"))
    _replace_text_in_doc(doc, "$vndId",      str(vendor.get("เลขผู้เสียภาษี", "-")))
    _replace_text_in_doc(doc, "$vndAddress", str(vendor.get("บ้านเลขที่", "-")))
    _replace_text_in_doc(doc, "$vndM",       str(vendor.get("หมู่", "-")))
    _replace_text_in_doc(doc, "$vndSt",      str(vendor.get("ถนน", "-")))
    _replace_text_in_doc(doc, "$vndZone",    str(vendor.get("แขวง/ตำบล", "-")))
    _replace_text_in_doc(doc, "$vndCity",    str(vendor.get("เขต/อำเภอ", "-")))
    _replace_text_in_doc(doc, "$vndProv",    str(vendor.get("จังหวัด", "-")))

    # ── หาตาราง item (ที่มี $desc[0]) ──
    item_table = None
    for t in doc.tables:
        text = "".join(cell.text for row in t.rows for cell in row.cells)
        if "$desc[0]" in text:
            item_table = t
            break

    if item_table is None:
        log("    ⚠️  ไม่พบตาราง item ใน template ใบสำคัญรับเงิน")
        return None

    # หา template row
    template_row_idx = None
    for idx, row in enumerate(item_table.rows):
        if "$desc[0]" in "".join(cell.text for cell in row.cells):
            template_row_idx = idx
            break

    if template_row_idx is None:
        return None

    # clone rows
    if len(slips) > 1:
        for _ in range(len(slips) - 1):
            _clone_row(item_table, template_row_idx)

    # เติมข้อมูลแต่ละ row
    total = 0
    for i, slip in enumerate(slips):
        tr      = item_table.rows[template_row_idx + i]._tr
        amt     = slip.get("amount") or 0
        amt_f   = int(amt)                              # บาท
        amt_st  = round((amt - amt_f) * 100)           # สตางค์
        desc    = slip.get("note") or slip.get("description") or "-"
        total  += amt

        _set_cell_text(tr, 0, desc)                    # รายการ
        _set_cell_text(tr, 1, f"{amt_f:,}")            # บาท
        _set_cell_text(tr, 2, f"{amt_st:02d}")         # สตางค์

    # รวม
    total_f  = int(total)
    total_st = round((total - total_f) * 100)

    _replace_text_in_doc(doc, "$intSumFTotal",  f"{total_f:,}")
    _replace_text_in_doc(doc, "$intSumStTotal", f"{total_st:02d}")
    _replace_text_in_doc(doc, "$thSumTotal",    baht_text(total))

    import tempfile
    tmp = Path(tempfile.mkdtemp()) / f"receipt_{month_name}_{day_str}.docx"
    doc.save(str(tmp))
    return tmp


def convert_to_pdf(docx_path: Path, output_dir: Path) -> Path | None:
    """แปลง docx เป็น PDF ด้วย LibreOffice พร้อม retry ถ้า lock"""
    import time as _time
    output_dir.mkdir(parents=True, exist_ok=True)

    for attempt in range(3):
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", str(output_dir), str(docx_path)],
                capture_output=True, text=True, timeout=120
            )
            pdf_name = docx_path.stem + ".pdf"
            pdf_path = output_dir / pdf_name
            if pdf_path.exists():
                return pdf_path

            # ถ้า convert ไม่สำเร็จ เช็ค lock file
            lock_dir = Path.home() / ".config/libreoffice/.~lock.*"
            lock_profile = Path.home() / ".config/libreoffice/user/.lock"
            for lock in [lock_profile, *Path.home().glob(".config/libreoffice/.~lock.*")]:
                if lock.exists():
                    lock.unlink(missing_ok=True)
                    log(f"    🔓 ลบ lock file: {lock}")

            if attempt < 2:
                log(f"    ⚠️  LibreOffice retry ({attempt + 1}/3)...")
                _time.sleep(3)
            else:
                log(f"    ⚠️  LibreOffice error: {result.stderr[:200]}")
        except subprocess.TimeoutExpired:
            log(f"    ⚠️  LibreOffice timeout — kill แล้ว retry...")
            subprocess.run(["pkill", "-f", "soffice"], capture_output=True)
            _time.sleep(3)
        except Exception as e:
            log(f"    ⚠️  Convert error: {e}")
            break
    return None

# ── Main ──────────────────────────────────────────────────────────────────────

def run(local_data_path: str | None = None) -> dict:
    data_root = Path(DATA_MOUNT)
    results   = {"new": 0, "skip": 0, "failed": 0, "monthly": {}}
    state     = load_state()

    # โหลด vendor list ครั้งเดียว
    vendors = load_vendors(force_reload=True)

    # ── local output dir สำหรับ PDF ทั้งหมด ──
    import tempfile
    local_output = Path(tempfile.mkdtemp()) / "output"
    local_output.mkdir(parents=True)

    # ถ้า sort ส่ง local_data_path มา ให้อ่านจากนั้น (เร็ว) แทนจาก mount
    if local_data_path:
        data_root = Path(local_data_path)

    # วน loop ปี/เดือน/วัน
    for year_dir in sorted(data_root.iterdir()):
        if not year_dir.is_dir() or year_dir.name in ("unclassified", "backups", "logs"):
            continue

        summary = load_summary(year_dir)

        for month_dir in sorted(year_dir.iterdir()):
            if not month_dir.is_dir():
                continue

            for day_dir in sorted(month_dir.iterdir()):
                if not day_dir.is_dir():
                    continue

                images_dir   = day_dir / "images"
                metadata_dir = day_dir / "metadata"

                # ไม่เช็คแค่ images_dir เพราะบางวัน metadata อยู่แต่รูปหาย (sync ไม่ครบในอดีต) —
                # gen_pdf.py อ่านแค่ metadata JSON ไม่ใช้รูปเลย ข้ามได้ก็ต่อเมื่อไม่มีทั้งคู่จริงๆ
                if not images_dir.exists() and not metadata_dir.exists():
                    continue

                # ── copy metadata มา local temp ก่อนอ่าน (เร็วกว่าอ่านจาก mount ทีละไฟล์) ──
                import tempfile, shutil as _shutil
                json_dir  = metadata_dir if metadata_dir.exists() else images_dir
                local_tmp = Path(tempfile.mkdtemp())
                try:
                    for jf in sorted(json_dir.glob("*.json")):
                        _shutil.copy(str(jf), str(local_tmp / jf.name))
                except Exception as e:
                    log(f"    ⚠️  copy metadata error: {e}")

                # หา slip JSON ทั้งหมดจาก local temp
                all_slips = []
                for jf in sorted(local_tmp.glob("*.json")):
                    data = json.loads(jf.read_text(encoding="utf-8"))
                    data["_json_path"] = str(json_dir / jf.name)  # path จริงบน mount
                    data["_local_json_path"] = str(jf)             # path local
                    all_slips.append(data)

                if not all_slips:
                    continue

                # ── แบ่ง slip ตาม route (note prefix) ──
                groups: dict[str, dict] = {}
                for s in all_slips:
                    note  = s.get("note") or ""
                    route = get_route(note)
                    sf    = route["subfolder"]
                    if sf not in groups:
                        groups[sf] = {
                            "slips": [],
                            "cert_template":    route["cert_template"],
                            "receipt_template": route["receipt_template"],
                        }
                    groups[sf]["slips"].append(s)

                for sf, group in groups.items():
                    group_slips      = group["slips"]
                    cert_tmpl_path   = group["cert_template"]
                    receipt_tmpl_path = group["receipt_template"]

                    docs_root      = local_output / year_dir.name / month_dir.name / day_dir.name / "docs"
                    local_docs     = docs_root / "ใบรับรองแทนใบเสร็จรับเงิน" / sf
                    local_receipt_dir = docs_root / "ใบสำคัญรับเงิน" / sf

                    new_slips = [
                        s for s in group_slips
                        if not is_generated(state, year_dir.name, month_dir.name,
                                            day_dir.name, sf, s)
                    ]

                    if not new_slips:
                        continue

                    log(f"  📄 {year_dir.name}/{month_dir.name}/{day_dir.name}/{sf} "
                        f"— slip ใหม่ {len(new_slips)} ใบ (รวม {len(group_slips)} ใบ)")

                    local_docs.mkdir(parents=True, exist_ok=True)

                    docx_path = fill_template(
                        group_slips, day_dir.name, month_dir.name,
                        year_dir.name, template_path=cert_tmpl_path,
                    )

                    if docx_path is None:
                        results["failed"] += 1
                        continue

                    date_prefix  = f"{year_dir.name}{MONTH_MAP_NUM.get(month_dir.name, '00')}{day_dir.name}"
                    # ใส่ sf ต่อท้ายชื่อไฟล์กัน collision ข้าม category (วันเดียวกันมีทั้ง uan/บุคคล ชื่อไฟล์ชนกันได้
                    # เพราะตอนนี้ทุก category ของเอกสารชนิดเดียวกันมารวมอยู่ใต้ folder แม่เดียวกัน)
                    final_name   = f"{date_prefix}-ใบรับรองแทนใบเสร็จรับเงิน-{sf}"
                    renamed_docx = local_docs / f"{final_name}.docx"
                    shutil.move(str(docx_path), str(renamed_docx))
                    pdf_path = convert_to_pdf(renamed_docx, local_docs)
                    renamed_docx.unlink(missing_ok=True)

                    if pdf_path is None:
                        log(f"    ❌ [{sf}] ใบรับรองฯ convert ล้มเหลว")
                        results["failed"] += 1
                        continue

                    log(f"    ✅ {final_name}.pdf")

                    from collections import defaultdict
                    slips_by_vendor: dict[str, list] = defaultdict(list)
                    for s in group_slips:
                        to_name = s.get("to_name") or s.get("to_account", "")
                        slips_by_vendor[to_name].append(s)

                    receipt_paths_rel = {}
                    for to_name, vendor_slips in slips_by_vendor.items():
                        vendor = find_vendor(to_name, vendors)
                        if vendor is None:
                            log(f"    ℹ️  ไม่เจอ vendor '{to_name}' — ข้ามใบสำคัญฯ")
                            continue
                        for s in vendor_slips:
                            s["vendor"] = vendor
                        safe_name = to_name.replace("/", "-").replace("\\", "-")
                        file_name = f"{date_prefix}-{safe_name}-{sf}"
                        local_receipt_dir.mkdir(parents=True, exist_ok=True)
                        receipt_docx = fill_template_receipt(
                            vendor_slips, day_dir.name, month_dir.name, year_dir.name,
                            template_path=receipt_tmpl_path,
                        )
                        if receipt_docx:
                            renamed_receipt = local_receipt_dir / f"{file_name}.docx"
                            shutil.move(str(receipt_docx), str(renamed_receipt))
                            receipt_pdf = convert_to_pdf(renamed_receipt, local_receipt_dir)
                            renamed_receipt.unlink(missing_ok=True)
                            if receipt_pdf:
                                log(f"    ✅ ใบสำคัญรับเงิน/{sf}/{file_name}.pdf")
                                receipt_paths_rel[to_name] = f"{file_name}.pdf"

                    mark_generated(state, year_dir.name, month_dir.name,
                                   day_dir.name, sf, group_slips, str(pdf_path))
                    save_state(state)

                    new_slip_paths = {s["_json_path"] for s in new_slips}
                    group_new = [s for s in group_slips if s["_json_path"] in new_slip_paths]
                    merge_to_summary(summary, month_dir.name, day_dir.name,
                                     group_new, pdf_path.name)

                    total_amt = sum(s.get("amount") or 0 for s in group_new)
                    results["monthly"][month_dir.name] = (
                        results["monthly"].get(month_dir.name, 0) + total_amt
                    )
                    log(f"    ✅ [{sf}] {pdf_path.name} (฿{total_amt:,.0f})")
                    results["new"] += 1

                    # เก็บข้อมูลไว้สำหรับบันทึก transactions หลัง sync เสร็จ (ต้องมีไฟล์บน Drive ก่อนหา link ได้)
                    results.setdefault("_pending_transactions", []).append({
                        "slips": group_new,
                        "category": sf,
                        "cert_filename": pdf_path.name,
                        "receipt_filenames": receipt_paths_rel,
                    })

        save_summary(year_dir, summary)

    log(f"\n✅ gen ใหม่: {results['new']}  ❌ ล้มเหลว: {results['failed']}")
    results["_local_output"] = str(local_output)
    return results


def sync_transactions_only(scope: str) -> dict:
    """
    บันทึก transactions ลง Sheets ใหม่จาก metadata ที่มีอยู่แล้ว โดย**ไม่ gen PDF ใหม่**
    (คำนวณชื่อไฟล์ cert/receipt ตาม naming convention เดิม แล้วให้ append_transactions()
    ไปหา Drive link เอาเอง — ใช้ตอน PDF/รูปอยู่บน Drive ครบแล้ว แค่ Sheets ไม่ update
    เช่น service account เพิ่งโดนลดสิทธิ์เป็น Viewer ตอน gen รอบก่อน)

    scope: "2026" / "2026/JUN" / "2026/JUN/15"
    """
    data_root = Path(DATA_MOUNT)
    scope_parts = scope.strip("/").split("/")
    scope_year  = scope_parts[0] if len(scope_parts) >= 1 else None
    scope_month = scope_parts[1] if len(scope_parts) >= 2 else None
    scope_day   = scope_parts[2] if len(scope_parts) >= 3 else None

    vendors = load_vendors(force_reload=True)
    from utils.transactions import append_transactions
    groups_done = 0

    for year_dir in sorted(data_root.iterdir()):
        if not year_dir.is_dir() or year_dir.name in ("unclassified", "backups", "logs"):
            continue
        if scope_year and year_dir.name != scope_year:
            continue

        for month_dir in sorted(year_dir.iterdir()):
            if not month_dir.is_dir():
                continue
            if scope_month and month_dir.name != scope_month:
                continue

            for day_dir in sorted(month_dir.iterdir()):
                if not day_dir.is_dir():
                    continue
                if scope_day and day_dir.name != scope_day:
                    continue

                images_dir   = day_dir / "images"
                metadata_dir = day_dir / "metadata"
                if not images_dir.exists() and not metadata_dir.exists():
                    continue

                json_dir = metadata_dir if metadata_dir.exists() else images_dir
                all_slips = []
                for jf in sorted(json_dir.glob("*.json")):
                    try:
                        all_slips.append(json.loads(jf.read_text(encoding="utf-8")))
                    except Exception as e:
                        log(f"    ⚠️  อ่าน {jf.name} ไม่ได้: {e}")

                if not all_slips:
                    continue

                groups: dict[str, list] = {}
                for s in all_slips:
                    note = s.get("note") or ""
                    sf   = get_route(note)["subfolder"]
                    groups.setdefault(sf, []).append(s)

                date_prefix = f"{year_dir.name}{MONTH_MAP_NUM.get(month_dir.name, '00')}{day_dir.name}"

                for sf, group_slips in groups.items():
                    cert_filename = f"{date_prefix}-ใบรับรองแทนใบเสร็จรับเงิน-{sf}.pdf"

                    from collections import defaultdict
                    slips_by_vendor: dict[str, list] = defaultdict(list)
                    for s in group_slips:
                        to_name = s.get("to_name") or s.get("to_account", "")
                        slips_by_vendor[to_name].append(s)

                    receipt_filenames = {}
                    for to_name, vendor_slips in slips_by_vendor.items():
                        vendor = find_vendor(to_name, vendors)
                        if vendor is None:
                            continue
                        for s in vendor_slips:
                            s["vendor"] = vendor
                        safe_name = to_name.replace("/", "-").replace("\\", "-")
                        receipt_filenames[to_name] = f"{date_prefix}-{safe_name}-{sf}.pdf"

                    log(f"   📦 {year_dir.name}/{month_dir.name}/{day_dir.name}/{sf} "
                        f"— {len(group_slips)} slips, cert={cert_filename}")
                    try:
                        append_transactions(
                            slips=group_slips,
                            category=sf,
                            cert_filename=cert_filename,
                            receipt_filenames=receipt_filenames,
                        )
                        groups_done += 1
                    except Exception as e:
                        log(f"   ⚠️  บันทึก transactions ไม่ได้: {e}")

    log(f"\n✅ บันทึก transactions {groups_done} groups")
    return {"groups": groups_done}


if __name__ == "__main__":
    import argparse, time, sys
    from run_pipeline import fmt_duration

    parser = argparse.ArgumentParser(description="Gen PDF ใบรับรองแทนใบเสร็จรับเงิน")
    parser.add_argument(
        "--regen",
        nargs=2,
        metavar=("SCOPE", "VALUE"),
        help="regen โดย bypass ตัวเช็ค เช่น --regen year 2026 | --regen month 2026/JUN | --regen day 2026/JUN/24"
    )
    parser.add_argument(
        "--transactions-only",
        metavar="SCOPE",
        help='บันทึก transactions ใหม่จาก metadata ที่มีอยู่แล้ว ไม่ gen PDF ใหม่ เช่น --transactions-only 2026/JUN'
    )
    args = parser.parse_args()

    from utils import notify

    if args.transactions_only:
        sync_transactions_only(args.transactions_only)
        raise SystemExit(0)

    # ── ไม่ได้ใส่ --regen มาเลย และเป็น terminal จริงๆ (ไม่ใช่ cron) → ถามเป็น wizard แทน ──
    if not args.regen and sys.stdin.isatty():
        ans = input("ต้องการ regen ไหม (y/N) > ").strip().lower()
        if ans == "y":
            year_ans = input("ระบุปี > ").strip()
            month_ans = input("ระบุเดือน (กรณีต้องการทั้งปี ใส่ 0) > ").strip()
            if not month_ans or month_ans == "0":
                args.regen = ("year", year_ans)
            else:
                month_name = MONTH_MAP.get(int(month_ans), month_ans) if month_ans.isdigit() else month_ans
                day_ans = input("ระบุวัน (กรณีต้องการทั้งเดือน ใส่ 0) > ").strip()
                if not day_ans or day_ans == "0":
                    args.regen = ("month", f"{year_ans}/{month_name}")
                else:
                    args.regen = ("day", f"{year_ans}/{month_name}/{int(day_ans):02d}")

    local_data_root = None
    local_data_path = None

    if args.regen:
        scope_type, scope_value = args.regen
        log(f"♻️  Regen mode: {scope_type} = {scope_value}")
        notify.send(f"🚀 Regen เริ่มแล้ว — scope: {scope_value}")

        from utils.state import load_state, save_state, reset_state
        state = load_state()
        count = reset_state(state, scope_value)
        save_state(state)
        log(f"   Reset {count} groups — กำลัง gen...\n")

        # ── local staging: copy แค่ folder "metadata/" ของ scope นี้มา local ก่อน (เร็วกว่าอ่านจาก
        # mount ทีละไฟล์) — ไม่เอา images/PDF เก่า/summary.json เพราะ gen_pdf.py ไม่ใช้เลย ──
        import tempfile
        local_data_root = Path(tempfile.mkdtemp()) / "regen_data"
        local_data_root.mkdir(parents=True, exist_ok=True)  # เผื่อ rclone ไม่เจอไฟล์เลยแล้วไม่สร้าง folder ให้
        log(f"📥 copy metadata scope={scope_value} มา local...")
        notify.send("📥 กำลัง copy metadata...")
        # ระบุ pattern ตรงๆ ทุกความลึกที่เป็นไปได้ (scope วัน/เดือน/ปี ทำให้ metadata/ อยู่ลึกไม่เท่ากัน)
        # ไม่พึ่ง "**/metadata/**" เพราะ rclone ไม่ match กรณี metadata อยู่ที่ root ตรงๆ (scope ระดับวัน)
        subprocess.run([
            "rclone", "copy", f"gdrive:SlipProcessor/data/{scope_value}",
            str(local_data_root / scope_value),
            "--include", "metadata/**",        # scope = วัน (metadata อยู่ root)
            "--include", "*/metadata/**",      # scope = เดือน (metadata อยู่ใต้ {day}/)
            "--include", "*/*/metadata/**",    # scope = ปี (metadata อยู่ใต้ {month}/{day}/)
            "--transfers", "16", "--checkers", "32", "--fast-list",
            "--stats", "30s", "-v",
            "--config", str(Path.home() / ".config/rclone/rclone.conf"),
        ])
        notify.send("✅ copy metadata เสร็จแล้ว")

        if not any(local_data_root.iterdir()):
            msg = f"⚠️ ไม่พบ metadata เลยสำหรับ scope={scope_value} — เช็คว่าวัน/เดือน/ปีถูกต้องไหม (ไม่มีอะไรให้ gen)"
            log(msg)
            notify.send(msg)
            raise SystemExit(0)

        local_data_path = str(local_data_root)

    t_total = time.time()

    log("\n── กำลัง gen PDF ──")
    notify.send("📄 กำลัง gen PDF...")
    t0 = time.time()
    result = run(local_data_path=local_data_path)
    gen_elapsed = fmt_duration(time.time() - t0)
    notify.send(f"✅ gen PDF เสร็จแล้ว ({gen_elapsed})")

    # ── sync local_output (PDF) ขึ้น Drive ──
    local_output = result.get("_local_output")
    pdf_ok = True
    sync_elapsed = "0 วิ"
    if local_output:
        log("\n── Sync PDFs ขึ้น Drive ──")
        notify.send("🔄 กำลัง sync PDF ขึ้น Drive...")
        t0 = time.time()
        from run_pipeline import sync_output_dir
        pdf_ok = sync_output_dir(local_output)
        sync_elapsed = fmt_duration(time.time() - t0)
        log("   ✅ PDFs synced" if pdf_ok else "   ❌ sync PDFs ล้มเหลว — ข้ามบันทึก transactions")
        notify.send(f"✅ sync PDF เสร็จแล้ว ({sync_elapsed})" if pdf_ok else "❌ sync PDF ล้มเหลว — ข้ามบันทึก transactions")

    # ── บันทึก transactions ลง Google Sheets (หลัง sync เสร็จ) ──
    pending = result.get("_pending_transactions", []) if pdf_ok else []
    t0 = time.time()
    if pending:
        log(f"\n── บันทึก Transactions ({len(pending)} groups) ──")
        notify.send(f"📊 กำลัง update transaction sheet ({len(pending)} groups)...")
        from utils.transactions import append_transactions
        for i, item in enumerate(pending, 1):
            log(f"   📦 [{i}/{len(pending)}] category={item['category']} cert={item['cert_filename']} receipts={list(item['receipt_filenames'].keys())}")
            try:
                append_transactions(
                    slips=item["slips"],
                    category=item["category"],
                    cert_filename=item["cert_filename"],
                    receipt_filenames=item["receipt_filenames"],
                )
            except Exception as e:
                log(f"   ⚠️  บันทึก transactions ไม่ได้: {e}")
    tx_elapsed = fmt_duration(time.time() - t0)
    total_elapsed = fmt_duration(time.time() - t_total)

    notify.send(
        f"✅ Pipeline เสร็จสิ้น (gen_pdf.py)\n"
        f"📄 gen ใหม่: {result.get('new', 0)}  ❌ ล้มเหลว: {result.get('failed', 0)}\n"
        f"⏱ Gen: {gen_elapsed}\n"
        f"🔄 Sync: {sync_elapsed}\n"
        f"📊 Transactions: {tx_elapsed} ({len(pending)} groups)\n"
        f"⏱ รวม: {total_elapsed}"
    )
    log(f"\n✅ เสร็จสิ้น — รวม {total_elapsed}")

    # ── cleanup temp files ──
    import shutil as _shutil
    if local_data_root:
        _shutil.rmtree(local_data_root, ignore_errors=True)
    if local_output:
        _shutil.rmtree(local_output, ignore_errors=True)
